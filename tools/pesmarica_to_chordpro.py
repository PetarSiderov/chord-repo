#!/usr/bin/env python3
"""
Pesmarica to ChordPro Converter
Scrapes song lyrics and chords from pesmarica.rs and converts to ChordPro format.
"""

import argparse
import re
import sys
import unicodedata
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup


def normalize_filename(text):
    """Normalize text for safe filename usage."""
    # Normalize unicode characters
    text = unicodedata.normalize('NFKD', text)
    # Remove or replace unsafe characters
    text = re.sub(r'[^\w\s\-]', '', text)
    # Replace spaces with underscores
    text = re.sub(r'[\s]+', '_', text)
    # Remove leading/trailing underscores
    text = text.strip('_')
    return text


def strip_tab_lines(lyrics):
    """Remove guitar tablature lines from lyrics."""
    lines = lyrics.split('\n')
    filtered_lines = []
    
    for line in lines:
        # Skip lines that look like guitar tabs (e.g., "e|--0--2--3--|")
        # TAB lines typically contain: string markers (e,B,G,D,A,E), pipes, dashes, numbers
        if re.match(r'^[eEbBgGdDaA]\|[\d\-\|\s]+\|?$', line.strip()):
            continue
        if re.match(r'^[eEbBgGdDaA]?\|[\-\d\|\s]+\|?$', line.strip()) and '--' in line:
            continue
        filtered_lines.append(line)
    
    return '\n'.join(filtered_lines)


def get_piano_hints(key):
    """Generate piano solo hints based on the key signature."""
    hints = {
        'C': 'Piano hints: C major scale (C D E F G A B). Try arpeggios: C-E-G, F-A-C, G-B-D.',
        'D': 'Piano hints: D major scale (D E F# G A B C#). Try arpeggios: D-F#-A, G-B-D, A-C#-E.',
        'E': 'Piano hints: E major scale (E F# G# A B C# D#). Try arpeggios: E-G#-B, A-C#-E, B-D#-F#.',
        'F': 'Piano hints: F major scale (F G A Bb C D E). Try arpeggios: F-A-C, Bb-D-F, C-E-G.',
        'G': 'Piano hints: G major scale (G A B C D E F#). Try arpeggios: G-B-D, C-E-G, D-F#-A.',
        'A': 'Piano hints: A major scale (A B C# D E F# G#). Try arpeggios: A-C#-E, D-F#-A, E-G#-B.',
        'B': 'Piano hints: B major scale (B C# D# E F# G# A#). Try arpeggios: B-D#-F#, E-G#-B, F#-A#-C#.',
        'Am': 'Piano hints: A minor scale (A B C D E F G). Try arpeggios: Am (A-C-E), Dm (D-F-A), Em (E-G-B).',
        'Bm': 'Piano hints: B minor scale (B C# D E F# G A). Try arpeggios: Bm (B-D-F#), Em (E-G-B), F#m (F#-A-C#).',
        'Cm': 'Piano hints: C minor scale (C D Eb F G Ab Bb). Try arpeggios: Cm (C-Eb-G), Fm (F-Ab-C), Gm (G-Bb-D).',
        'Dm': 'Piano hints: D minor scale (D E F G A Bb C). Try arpeggios: Dm (D-F-A), Gm (G-Bb-D), Am (A-C-E).',
        'Em': 'Piano hints: E minor scale (E F# G A B C D). Try arpeggios: Em (E-G-B), Am (A-C-E), Bm (B-D-F#).',
        'Fm': 'Piano hints: F minor scale (F G Ab Bb C Db Eb). Try arpeggios: Fm (F-Ab-C), Bbm (Bb-Db-F), Cm (C-Eb-G).',
        'Gm': 'Piano hints: G minor scale (G A Bb C D Eb F). Try arpeggios: Gm (G-Bb-D), Cm (C-Eb-G), Dm (D-F-A).',
        'F#m': 'Piano hints: F# minor scale (F# G# A B C# D E). Try arpeggios: F#m (F#-A-C#), Bm (B-D-F#), C#m (C#-E-G#).',
        'G#m': 'Piano hints: G# minor scale (G# A# B C# D# E F#). Try arpeggios: G#m (G#-B-D#), C#m (C#-E-G#), D#m (D#-F#-A#).',
        'C#m': 'Piano hints: C# minor scale (C# D# E F# G# A B). Try arpeggios: C#m (C#-E-G#), F#m (F#-A-C#), G#m (G#-B-D#).',
    }
    return hints.get(key, '')


def scrape_pesmarica(url):
    """Scrape song lyrics and chords from pesmarica.rs."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        response.encoding = 'utf-8'
    except requests.RequestException as e:
        print(f"Error fetching URL: {e}", file=sys.stderr)
        return None
    
    soup = BeautifulSoup(response.text, 'html.parser')
    
    # Find the song content - pesmarica.rs typically uses div with class 'akordi'
    content_div = soup.find('div', class_='akordi')
    if not content_div:
        # Try alternative selectors
        content_div = soup.find('pre', class_='akordi')
    if not content_div:
        content_div = soup.find('div', id='akordi')
    
    if not content_div:
        print("Could not find lyrics content on the page", file=sys.stderr)
        return None
    
    # Extract text content preserving line breaks
    lyrics = content_div.get_text(separator='\n', strip=False)
    
    return lyrics


def convert_to_chordpro(lyrics, title, artist, key=None, tempo=None, meter=None, 
                       url=None, strip_tabs=True, piano_hints=False):
    """Convert lyrics with chords to ChordPro format."""
    
    # Strip TAB lines if requested
    if strip_tabs:
        lyrics = strip_tab_lines(lyrics)
    
    # Start building ChordPro content
    chordpro = []
    
    # Add metadata directives
    if title:
        chordpro.append(f'{{title: {title}}}')
    if artist:
        chordpro.append(f'{{artist: {artist}}}')
    if key:
        chordpro.append(f'{{key: {key}}}')
    if tempo:
        chordpro.append(f'{{tempo: {tempo}}}')
    if meter:
        chordpro.append(f'{{time: {meter}}}')
    
    # Add source URL as comment
    if url:
        chordpro.append(f'{{comment: Source: {url}}}')
    
    chordpro.append('')  # Empty line after metadata
    
    # Process lyrics line by line
    # Chords are typically on separate lines above lyrics in pesmarica.rs
    lines = lyrics.split('\n')
    
    for i, line in enumerate(lines):
        # Check if this line contains chords (uppercase letters, possibly with m, 7, sus, etc.)
        # Typical chord line pattern: starts with uppercase, contains chord symbols
        if line.strip() and re.match(r'^[A-G]', line.strip()):
            # Check if it's likely a chord line (contains typical chord symbols)
            if re.search(r'[A-G][#b]?(m|maj|min|sus|aug|dim|add)?\d*(/[A-G][#b]?)?', line):
                # This might be a chord line
                # In ChordPro, we put chords inline with lyrics using [chord]
                # For now, keep chord lines as-is with ChordPro chord notation
                # We'll wrap them in brackets
                chord_line = line
                # Replace chord patterns with [chord] notation
                chord_line = re.sub(r'\b([A-G][#b]?(m|maj|min|sus|aug|dim|add)?\d*(/[A-G][#b]?)?)\b', 
                                   r'[\1]', chord_line)
                chordpro.append(chord_line)
                continue
        
        # Regular lyric line or empty line
        chordpro.append(line)
    
    # Add piano hints if requested
    if piano_hints and key:
        hints = get_piano_hints(key)
        if hints:
            chordpro.append('')
            chordpro.append(f'{{comment: {hints}}}')
    
    return '\n'.join(chordpro)


def save_to_file(content, filename, output_dir='lyrics-output'):
    """Save content to a file in the output directory."""
    import os
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    filepath = os.path.join(output_dir, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return filepath


def create_docx(chordpro_content, filename, output_dir='lyrics-output'):
    """Create a DOCX file from ChordPro content (optional feature)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Skipping DOCX generation.", file=sys.stderr)
        print("Install with: pip install python-docx", file=sys.stderr)
        return None
    
    import os
    
    doc = Document()
    
    # Parse ChordPro content
    lines = chordpro_content.split('\n')
    
    # Add title and metadata at the top
    for line in lines:
        if line.startswith('{title:'):
            title = line.replace('{title:', '').replace('}', '').strip()
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('{artist:'):
            artist = line.replace('{artist:', '').replace('}', '').strip()
            p = doc.add_paragraph(artist)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.italic = True
        elif line.startswith('{key:') or line.startswith('{tempo:') or line.startswith('{time:'):
            text = line.replace('{', '').replace('}', '').strip()
            p = doc.add_paragraph(text)
            p.runs[0].font.size = Pt(10)
        elif line.startswith('{comment:'):
            comment = line.replace('{comment:', '').replace('}', '').strip()
            p = doc.add_paragraph(comment)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        elif not line.startswith('{'):
            # Regular content line
            if line.strip():
                # Check if line contains chords [chord]
                if '[' in line and ']' in line:
                    # Line with chords - make chords bold and colored
                    p = doc.add_paragraph()
                    parts = re.split(r'(\[[^\]]+\])', line)
                    for part in parts:
                        if part.startswith('[') and part.endswith(']'):
                            # This is a chord
                            chord_text = part[1:-1]  # Remove brackets
                            run = p.add_run(chord_text + ' ')
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 255)
                        else:
                            # Regular text
                            p.add_run(part)
                else:
                    # Regular text line
                    doc.add_paragraph(line)
            else:
                # Empty line
                doc.add_paragraph()
    
    # Save the document
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filepath


def main():
    parser = argparse.ArgumentParser(
        description='Convert song from pesmarica.rs to ChordPro format'
    )
    parser.add_argument('--url', required=True, help='URL of the song on pesmarica.rs')
    parser.add_argument('--title', required=True, help='Song title')
    parser.add_argument('--artist', required=True, help='Artist name')
    parser.add_argument('--key', help='Song key (e.g., G, Am, F#m)')
    parser.add_argument('--tempo', help='Tempo in BPM')
    parser.add_argument('--meter', help='Time signature (e.g., 4/4, 3/4)')
    parser.add_argument('--docx', action='store_true', help='Generate DOCX file')
    parser.add_argument('--keep-tabs', action='store_true', 
                       help='Keep guitar TAB lines (default: strip them)')
    parser.add_argument('--strip-tabs', action='store_true', dest='strip_tabs',
                       help='Strip guitar TAB lines (default behavior)')
    parser.add_argument('--piano-hints', action='store_true',
                       help='Add piano solo hints based on key')
    parser.add_argument('--output-dir', default='lyrics-output',
                       help='Output directory (default: lyrics-output)')
    
    args = parser.parse_args()
    
    # Determine if we should strip tabs (default is True unless --keep-tabs is set)
    strip_tabs = not args.keep_tabs
    
    print(f"Fetching song from: {args.url}")
    lyrics = scrape_pesmarica(args.url)
    
    if not lyrics:
        print("Failed to scrape lyrics", file=sys.stderr)
        return 1
    
    print(f"Converting to ChordPro format...")
    chordpro = convert_to_chordpro(
        lyrics, 
        args.title, 
        args.artist, 
        args.key, 
        args.tempo, 
        args.meter,
        args.url,
        strip_tabs=strip_tabs,
        piano_hints=args.piano_hints
    )
    
    # Generate safe filename
    safe_title = normalize_filename(args.title)
    safe_artist = normalize_filename(args.artist)
    base_filename = f"{safe_artist}_{safe_title}"
    
    # Save ChordPro file
    pro_filename = f"{base_filename}.pro"
    pro_path = save_to_file(chordpro, pro_filename, args.output_dir)
    print(f"Saved ChordPro file: {pro_path}")
    
    # Generate DOCX if requested
    if args.docx:
        docx_filename = f"{base_filename}.docx"
        docx_path = create_docx(chordpro, docx_filename, args.output_dir)
        if docx_path:
            print(f"Saved DOCX file: {docx_path}")
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
